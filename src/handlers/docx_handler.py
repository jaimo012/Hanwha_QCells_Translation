"""
Word 문서 핸들러 (DOCX Handler)

Word 문서(.docx)의 번역을 처리합니다.
- 본문, 표, 텍스트 상자를 모두 탐색
- 서식(폰트, 정렬, 들여쓰기 등)을 완벽하게 보존
- 토큰 사용량 추적
"""

import os
import time
import subprocess
from docx import Document
from docx.text.paragraph import Paragraph

from ..config import BATCH_SIZE_DOCX, API_DELAY_SECONDS, AUTO_SAVE_INTERVAL
from ..utils import has_korean
from ..translator import translate_batch


# 파일 저장 재시도 설정
SAVE_MAX_RETRIES = 5
SAVE_RETRY_DELAY = 3  # 초


def kill_word_processes():
    """
    실행 중인 Word 프로세스를 강제 종료합니다.
    파일이 잠겨있을 때 사용합니다.
    """
    try:
        subprocess.run(
            ['taskkill', '/F', '/IM', 'WINWORD.EXE'],
            capture_output=True,
            timeout=10
        )
        time.sleep(2)  # Word 프로세스 종료 대기
    except Exception:
        pass


def save_document_with_retry(doc, file_path, max_retries=SAVE_MAX_RETRIES):
    """
    문서를 저장하며, 실패 시 재시도합니다.
    
    Args:
        doc: python-docx Document 객체
        file_path (str): 저장할 파일 경로
        max_retries (int): 최대 재시도 횟수
        
    Returns:
        bool: 저장 성공 여부
        
    Raises:
        Exception: 모든 재시도 실패 시
    """
    for attempt in range(1, max_retries + 1):
        try:
            doc.save(file_path)
            return True
        except PermissionError as e:
            if attempt < max_retries:
                print(f"\n   ⚠️ 파일 저장 실패 (시도 {attempt}/{max_retries}): 파일이 사용 중입니다")
                print(f"   🔄 {SAVE_RETRY_DELAY}초 후 재시도...")
                
                # Word 프로세스 강제 종료 시도
                if attempt >= 2:
                    print(f"   🔧 Word 프로세스 종료 시도...")
                    kill_word_processes()
                else:
                    time.sleep(SAVE_RETRY_DELAY)
            else:
                raise Exception(f"파일 저장 실패 (모든 재시도 실패): {e}")
        except Exception as e:
            if attempt < max_retries:
                print(f"\n   ⚠️ 파일 저장 오류 (시도 {attempt}/{max_retries}): {e}")
                time.sleep(SAVE_RETRY_DELAY)
            else:
                raise
    
    return False


def iter_docx_paragraphs(doc):
    """
    Word 문서의 모든 문단을 순회하는 제너레이터입니다.
    
    본문, 표, 도형(텍스트 상자)을 통합하여 순회합니다.
    
    Args:
        doc: python-docx Document 객체
        
    Yields:
        Paragraph: 문서 내의 각 문단 객체
    """
    # 1. 본문 (Body)
    for p in doc.paragraphs:
        yield p

    # 2. 표 (Tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    
    # 3. 도형/텍스트 상자 (Shapes) - XML 파싱
    for element in doc.element.body.iter():
        if element.tag.endswith('txbxContent'):
            for child in element.iter():
                if child.tag.endswith('p'):
                    yield Paragraph(child, doc)


def copy_docx_style_and_replace(paragraph, new_text):
    """
    텍스트를 교체하면서 문단의 모든 서식을 보존합니다.
    
    Args:
        paragraph: 대상 문단 객체
        new_text (str): 교체할 새 텍스트
    """
    # 1. 문단 속성(Paragraph Format) 백업
    p_format = paragraph.paragraph_format
    
    para_backup = {
        'alignment': paragraph.alignment,
        'left_indent': p_format.left_indent,
        'right_indent': p_format.right_indent,
        'first_line_indent': p_format.first_line_indent,
        'space_before': p_format.space_before,
        'space_after': p_format.space_after,
        'line_spacing': p_format.line_spacing,
        'line_spacing_rule': p_format.line_spacing_rule,
        'keep_together': p_format.keep_together,
        'keep_with_next': p_format.keep_with_next
    }

    # 2. 글자 속성(Run Style) 백업
    run_backup = {}
    if paragraph.runs:
        ref_run = paragraph.runs[0]
        run_backup = {
            'font_name': ref_run.font.name,
            'font_size': ref_run.font.size,
            'bold': ref_run.bold,
            'italic': ref_run.italic,
            'underline': ref_run.underline,
            'color': ref_run.font.color.rgb if ref_run.font.color else None,
            'style': ref_run.style
        }

    # 3. 텍스트 교체
    paragraph.text = new_text

    # 4. 문단 속성 복구
    paragraph.alignment = para_backup['alignment']
    new_p_format = paragraph.paragraph_format
    
    new_p_format.left_indent = para_backup['left_indent']
    new_p_format.right_indent = para_backup['right_indent']
    new_p_format.first_line_indent = para_backup['first_line_indent']
    new_p_format.space_before = para_backup['space_before']
    new_p_format.space_after = para_backup['space_after']
    new_p_format.line_spacing = para_backup['line_spacing']
    new_p_format.line_spacing_rule = para_backup['line_spacing_rule']
    new_p_format.keep_together = para_backup['keep_together']
    new_p_format.keep_with_next = para_backup['keep_with_next']

    # 5. 글자 속성 복구
    if paragraph.runs and run_backup:
        new_run = paragraph.runs[0]
        new_run.bold = run_backup.get('bold')
        new_run.italic = run_backup.get('italic')
        new_run.underline = run_backup.get('underline')
        new_run.style = run_backup.get('style')
        
        if run_backup.get('font_name'):
            new_run.font.name = run_backup['font_name']
        if run_backup.get('font_size'):
            new_run.font.size = run_backup['font_size']
        if run_backup.get('color'):
            try:
                new_run.font.color.rgb = run_backup['color']
            except:
                pass


def process_docx(file_path, context, sheets_manager=None, row_index=None):
    """
    Word 문서를 번역합니다.
    
    Args:
        file_path (str): 원본 Word 파일 경로
        context (str): 번역 지침 (Context)
        sheets_manager (SheetsManager, optional): 시트 관리자 (토큰 추적용)
        row_index (int, optional): 시트 행 번호
        
    Returns:
        str: 번역된 파일의 경로 (성공 시)
        None: 실패 시
    """
    print(f"📖 Word 처리 중: {os.path.basename(file_path)}")
    
    doc = Document(file_path)
    # 작업 파일을 그대로 덮어쓰기 (main.py에서 이미 " - en" 파일 생성)
    new_path = file_path
    
    batch_queue = []
    total_count = 0
    batch_cycle = 0
    total_input_tokens = 0
    total_output_tokens = 0
    
    for paragraph in iter_docx_paragraphs(doc):
        # paragraph.text가 None일 수 있으므로 안전하게 처리
        text = paragraph.text
        if text is None:
            continue
        text = text.strip()
        
        if text and has_korean(text):
            batch_queue.append(paragraph)
            
            if len(batch_queue) >= BATCH_SIZE_DOCX:
                texts = [p.text for p in batch_queue]
                translated, input_tokens, output_tokens = translate_batch(texts, context)
                
                total_input_tokens += input_tokens
                total_output_tokens += output_tokens
                
                if len(translated) == len(batch_queue):
                    for obj, trans_text in zip(batch_queue, translated):
                        copy_docx_style_and_replace(obj, trans_text)
                    total_count += len(translated)
                    batch_cycle += 1
                    
                    # 실시간 진행 상황 (배치 횟수만 표시)
                    print(f"   ▶ 배치 {batch_cycle}회 진행 중...          ", end="\r")

                    # 중간 저장
                    if batch_cycle % AUTO_SAVE_INTERVAL == 0:
                        print()  # 줄바꿈
                        print(f"   💾 [자동저장] 데이터 보호를 위해 중간 저장 중...")
                        save_document_with_retry(doc, new_path)
                        
                        # 시트에 토큰 사용량 업데이트
                        if sheets_manager and row_index:
                            sheets_manager.update_tokens(
                                row_index,
                                total_input_tokens,
                                total_output_tokens
                            )
                            total_input_tokens = 0
                            total_output_tokens = 0
                
                batch_queue = []
                time.sleep(API_DELAY_SECONDS)

    # 잔여 처리 (배치 크기보다 적은 남은 데이터)
    if batch_queue:
        print(f"\n   🔄 잔여 {len(batch_queue)}개 처리 중...")
        texts = [p.text for p in batch_queue]
        translated, input_tokens, output_tokens = translate_batch(texts, context)
        
        total_input_tokens += input_tokens
        total_output_tokens += output_tokens
        
        if len(translated) == len(batch_queue):
            for obj, trans_text in zip(batch_queue, translated):
                copy_docx_style_and_replace(obj, trans_text)
            total_count += len(translated)
            batch_cycle += 1
        print(f"   ✅ 잔여 처리 완료")

    print(f"\n   💾 최종 저장 중...")
    save_document_with_retry(doc, new_path)
    print(f"   ✅ 파일 저장 완료")
    
    # 최종 토큰 사용량 업데이트
    if sheets_manager and row_index:
        if total_input_tokens > 0 or total_output_tokens > 0:
            sheets_manager.update_tokens(
                row_index,
                total_input_tokens,
                total_output_tokens
            )
    
    print()  # 진행 상황 줄 종료
    print(f"   ✅ Word 번역 완료: {batch_cycle}개 배치, {total_count}개 문장")
    return new_path
